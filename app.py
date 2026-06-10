import sqlite3, webbrowser, os, platform, subprocess, shutil, zipfile, re, hashlib, secrets
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import load_workbook, Workbook
import qrcode

APP='GST Billing Demo'
DEFAULT_PASSWORD='1234567890'
BASE=Path(__file__).resolve().parent
DATA=BASE/'data'; ASSETS=BASE/'assets'; MASTER=DATA/'master'
DB=DATA/'demo_billing.db'; LOGO=ASSETS/'full-logo.png'
for p in [DATA, MASTER, ASSETS]: p.mkdir(parents=True, exist_ok=True)

BUSINESS={
    'name':'Demo Textile Research Lab',
    'gstin':'00ABCDE1234A1Z0',
    'address':'123 Demo Industrial Area, Sample City - 000000',
    'email':'demo@example.com',
    'phone':'99999-99999',
    'state':'Rajasthan',
    'code':'08',
    'bank':'DEMO BANK, Main Branch, Sample City',
    'account':'00000000000000',
    'ifsc':'DEMO0001234',
    'upi':''
}

VALID_STATE_CODES={
    '01','02','03','04','05','06','07','08','09','10','11','12','13','14','15','16','17','18','19','20','21','22','23','24','25','26','27','28','29','30','31','32','33','34','35','36','37','38','97'
}
GSTIN_RE=re.compile(r'^[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z]$')

def hash_password(password, salt=None):
    salt = salt or secrets.token_hex(16)
    digest = hashlib.sha256((salt + password).encode('utf-8')).hexdigest()
    return f'sha256${salt}${digest}'

def check_password(password, stored):
    stored = stored or ''
    if stored.startswith('sha256$'):
        try:
            _, salt, digest = stored.split('$', 2)
            return hashlib.sha256((salt + password).encode('utf-8')).hexdigest() == digest
        except Exception:
            return False
    return password == stored

def safe_now():
    now=datetime.now()
    # If a PC clock is obviously wrong, do not auto-fill dates/timestamps with it.
    if now.year < 2024 or now.year > datetime.now().year + 10:
        return None
    return now

def safe_date_default():
    now=safe_now()
    return now.strftime('%d/%m/%Y') if now else ''

def safe_timestamp():
    now=safe_now() or datetime(2026,1,1,0,0,0)
    return now.strftime('%Y%m%d_%H%M%S')

def validate_gstin(gstin):
    gstin=str(gstin or '').strip().upper()
    return bool(GSTIN_RE.match(gstin))

def normalize_bill_no(bill_no):
    raw=str(bill_no or '').strip()
    if not raw.isdigit():
        raise ValueError('Bill number must contain only digits. Example: 0001')
    n=int(raw)
    if n <= 0 or n > 9999:
        raise ValueError('Bill number must be between 0001 and 9999.')
    return f'{n:04d}'

def state_code_from_gstin(gstin):
    g=str(gstin or '').strip().upper()
    return g[:2] if len(g) >= 2 and g[:2].isdigit() else ''

def validate_invoice_date_str(date_text):
    dt=parse_invoice_date(date_text) if 'parse_invoice_date' in globals() else None
    if not dt:
        return None, 'Date must be in dd/mm/yyyy format.'
    current_year=datetime.now().year
    max_year=current_year + 1 if 2024 <= current_year <= 2090 else 2035
    if dt.year < 2020 or dt.year > max_year:
        return None, f'Invoice year looks wrong. Please use a year between 2020 and {max_year}.'
    return dt, ''

def money(x):
    try:
        if x in [None,'','-']: return 0.0
        return round(float(str(x).replace(',','').strip()),2)
    except Exception:
        return 0.0

def ensure_master_files():
    DATA.mkdir(parents=True, exist_ok=True)
    MASTER.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    customers=MASTER/'customers.xlsx'
    if not customers.exists():
        wb=Workbook(); ws=wb.active; ws.title='Customers'
        ws.append(['name','address','gstin','state','state code','phone'])
        ws.append(['RAM KUMAR TEXTILE PVT. LTD.','BHILWARA','08AABCR8798J1ZK','Rajasthan','08',''])
        wb.save(customers)
    rates=MASTER/'rate_list.xlsx'
    if not rates.exists():
        wb=Workbook(); ws=wb.active; ws.title='Rate List'
        ws.append(['test','operation/report','hsn code','quantity','amount'])
        ws.append(['FABRIC DEFECT ANALYSIS','REPORT NO- DEMO/302','998346','1','3000'])
        wb.save(rates)

def db():
    DATA.mkdir(parents=True, exist_ok=True); ensure_master_files()
    con=sqlite3.connect(DB); cur=con.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS invoices(
        id INTEGER PRIMARY KEY AUTOINCREMENT, bill_no TEXT UNIQUE, date TEXT, party TEXT, gstin TEXT,
        taxable REAL, discount REAL, freight REAL, cgst REAL, sgst REAL, igst REAL, total REAL,
        docx_path TEXT, created_at TEXT)''')
    cur.execute('''CREATE TABLE IF NOT EXISTS settings(key TEXT PRIMARY KEY, value TEXT)''')
    cur.execute('''CREATE TABLE IF NOT EXISTS purchases(
        id INTEGER PRIMARY KEY AUTOINCREMENT, bill_no TEXT, date TEXT, party TEXT, gstin TEXT,
        taxable REAL, cgst REAL, sgst REAL, igst REAL, total REAL, note TEXT, created_at TEXT)''')
    cur.execute('INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)',('password_hash', hash_password(DEFAULT_PASSWORD)))
    # Migrate older plain-text password setting into password_hash once.
    cur.execute('SELECT value FROM settings WHERE key=?', ('password',))
    old_pw = cur.fetchone()
    cur.execute('SELECT value FROM settings WHERE key=?', ('password_hash',))
    current_hash = cur.fetchone()
    if old_pw and old_pw[0] and (not current_hash or not str(current_hash[0]).startswith('sha256$')):
        cur.execute('REPLACE INTO settings(key,value) VALUES(?,?)', ('password_hash', hash_password(str(old_pw[0]))))
    cur.execute('INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)',('records_folder', str(DATA)))
    for k,v in BUSINESS.items():
        cur.execute('INSERT OR IGNORE INTO settings(key,value) VALUES(?,?)',(k,v))
    con.commit(); return con

def get_settings():
    con=db(); cur=con.cursor(); cur.execute('SELECT key,value FROM settings'); d=dict(cur.fetchall()); con.close(); return d

def set_setting(k,v):
    con=db(); con.execute('REPLACE INTO settings(key,value) VALUES(?,?)',(k,v)); con.commit(); con.close()

def records_base():
    """Master folder where office records are saved. Default stays inside app data.
    User can change it from Settings without Python installed."""
    try:
        s=get_settings()
        raw=(s.get('records_folder') or '').strip()
        base=Path(raw) if raw else DATA
    except Exception:
        base=DATA
    base.mkdir(parents=True, exist_ok=True)
    for name in ['Invoices','GST Exports','Monthly Reports','Purchases','Backups','QR Temp']:
        (base/name).mkdir(parents=True, exist_ok=True)
    return base

def records_path(folder_name):
    path=records_base()/folder_name
    path.mkdir(parents=True, exist_ok=True)
    return path

def fiscal_year(dt):
    return f'{dt.year-1}-{str(dt.year)[-2:]}' if dt.month < 4 else f'{dt.year}-{str(dt.year+1)[-2:]}'

def rows_from_xlsx(path):
    if not path.exists(): return []
    wb=load_workbook(path, data_only=True); ws=wb.active
    headers=[str(c.value or '').strip().lower() for c in ws[1]]
    out=[]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row): continue
        out.append({headers[i]: row[i] if i<len(row) else '' for i in range(len(headers))})
    return out

def load_customers(): return rows_from_xlsx(MASTER/'customers.xlsx')
def load_rates(): return rows_from_xlsx(MASTER/'rate_list.xlsx')

def add_customer_to_file(cust):
    path=MASTER/'customers.xlsx'
    name=str(cust.get('name','')).strip()
    gstin=str(cust.get('gstin','')).strip().upper()
    if not name:
        raise ValueError('Customer name is required.')
    if gstin and not validate_gstin(gstin):
        raise ValueError('Customer GSTIN format is invalid.')
    existing=load_customers()
    for c in existing:
        if str(c.get('name','')).strip().lower()==name.lower() or (gstin and str(c.get('gstin','')).strip().upper()==gstin):
            raise ValueError('This customer already exists in customer list.')
    try:
        wb=load_workbook(path); ws=wb.active
        ws.append([name, cust.get('address',''), gstin, cust.get('state','Rajasthan'), cust.get('state code','08'), cust.get('phone','')])
        wb.save(path)
    except PermissionError:
        raise PermissionError('customers.xlsx is open in Excel. Please close it and try again.')

def shade(cell, color):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tcPr.append(shd)

def set_border(cell):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right'):
        tag='w:'+edge; elem=borders.find(qn(tag))
        if elem is None:
            elem=OxmlElement(tag); borders.append(elem)
        elem.set(qn('w:val'),'single'); elem.set(qn('w:sz'),'6'); elem.set(qn('w:color'),'555555')

def set_cell_text(cell, text, bold=False, size=9, align=None):
    cell.text=''
    p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(str(text))
    r.bold=bold; r.font.size=Pt(size)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_border(cell)

def make_qr(text, fname):
    img=qrcode.make(text)
    path=records_path('QR Temp')/f'{fname}.png'; img.save(path); return path

def amount_words_simple(total):
    """Convert integer rupees to Indian English words for invoices."""
    n=int(round(money(total)))
    if n==0:
        return 'RUPEES ZERO ONLY'
    ones=['','ONE','TWO','THREE','FOUR','FIVE','SIX','SEVEN','EIGHT','NINE','TEN','ELEVEN','TWELVE','THIRTEEN','FOURTEEN','FIFTEEN','SIXTEEN','SEVENTEEN','EIGHTEEN','NINETEEN']
    tens=['','','TWENTY','THIRTY','FORTY','FIFTY','SIXTY','SEVENTY','EIGHTY','NINETY']
    def under_100(num):
        if num<20: return ones[num]
        return tens[num//10] + ((' ' + ones[num%10]) if num%10 else '')
    def under_1000(num):
        if num<100: return under_100(num)
        rest=num%100
        return ones[num//100] + ' HUNDRED' + ((' ' + under_100(rest)) if rest else '')
    parts=[]
    crore=n//10000000; n%=10000000
    lakh=n//100000; n%=100000
    thousand=n//1000; n%=1000
    if crore: parts.append(under_1000(crore)+' CRORE')
    if lakh: parts.append(under_1000(lakh)+' LAKH')
    if thousand: parts.append(under_1000(thousand)+' THOUSAND')
    if n: parts.append(under_1000(n))
    return 'RUPEES ' + ' '.join(parts).strip() + ' ONLY'

def invoice_docx(values):
    s=get_settings()
    inv_date=datetime.strptime(values['date'],'%d/%m/%Y')
    fy=fiscal_year(inv_date); month=inv_date.strftime('%B').upper()
    outdir=records_path('Invoices')/fy/month; outdir.mkdir(parents=True, exist_ok=True)
    bill=normalize_bill_no(values['bill_no'])
    docx_path=outdir/f'GST INVOICE {bill}.docx'
    total=money(values['total'])

    upi=s.get('upi','').strip()
    pay_qr=None
    if upi:
        pay_link=f'upi://pay?pa={upi}&pn={s.get("name","").replace(" ", "%20")}&am={total:.2f}&cu=INR&tn=Invoice%20{bill}'
        pay_qr=make_qr(pay_link, f'pay_{bill}')

    doc=Document();
    style=doc.styles['Normal']; style.font.name='Cambria'; style.font.size=Pt(10)
    sec=doc.sections[0]
    sec.top_margin=Inches(0.32); sec.bottom_margin=Inches(0.32); sec.left_margin=Inches(0.42); sec.right_margin=Inches(0.42)

    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run('TAX INVOICE'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=RGBColor(0,0,0)

    t=doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    left,right=t.rows[0].cells
    if LOGO.exists(): left.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.35))
    else: left.text=s.get('name','DEMO')
    p=right.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run(s.get('name','Demo Textile Research Lab')+'\n'); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(0,0,0)
    # Keep contact lines separate so phone/email do not overlap in Word
    p.add_run(f'GSTIN : {s.get("gstin","")}\n')
    p.add_run(f'{s.get("address","")}\n')
    p.add_run(f'Email : {s.get("email","")}\n')
    p.add_run(f'Mobile No. : {s.get("phone","")}')
    for c in (left,right): set_border(c)

    info=doc.add_table(rows=5, cols=4); info.alignment=WD_TABLE_ALIGNMENT.CENTER
    labels=[('Details of Receiver (Billed to)','', 'GST Invoice No.:',bill),('Name :',values['party'],'Invoice Date:',values['date']),('Address :',values['address'],'Despatch Through :','-'),('Mob. No.-',values.get('phone',''),'Reverse Charge(Y/N):','-'),('GSTIN :',values['gstin'], f'State : {values["state"]}', f'Code : {values["code"]}')]
    for i,row in enumerate(labels):
        for j,val in enumerate(row): set_cell_text(info.cell(i,j), val, bold=(j in [0,2] or i==0), size=9)
    # Plain black-and-white invoice: no grey background shading.

    items=doc.add_table(rows=1+len(values['items']), cols=7); items.alignment=WD_TABLE_ALIGNMENT.CENTER
    heads=['S.N.','Tests','Report No.','HSN Code','Quantity','Amount','Total']
    for j,h in enumerate(heads):
        set_cell_text(items.cell(0,j),h,bold=True,size=9,align=WD_ALIGN_PARAGRAPH.CENTER)
    for i,it in enumerate(values['items'], start=1):
        row=items.rows[i].cells
        report_no=str(it.get('operation','') or '').strip()
        # User types the report number; print exactly that value in a separate column.
        vals=[str(i), it.get('test',''), report_no, it.get('hsn',''), it.get('qty',''), f'{money(it.get("amount")):.2f}', f'{money(it.get("line_total")):.2f}']
        for j,v in enumerate(vals):
            set_cell_text(row[j],v,size=9,align=WD_ALIGN_PARAGRAPH.CENTER if j not in [1,2] else None)

    calc=doc.add_table(rows=9, cols=2); calc.alignment=WD_TABLE_ALIGNMENT.RIGHT
    lines=[('Amount',values['gross']),('Discount',values['discount'] if money(values['discount']) else '-'),('Freight',values['freight'] if money(values['freight']) else '-'),('Total Amt. Before Tax',values['taxable']),('Add: SGST 9%',values['sgst'] if money(values['sgst']) else '-'),('Add: CGST 9%',values['cgst'] if money(values['cgst']) else '-'),('Add: IGST',values['igst'] if money(values['igst']) else '-'),('Total Tax Amt.', money(values['cgst'])+money(values['sgst'])+money(values['igst'])),('Total Amt. After Tax',values['total'])]
    for i,(a,b) in enumerate(lines):
        set_cell_text(calc.cell(i,0),a,bold=i in [3,8],size=9)
        set_cell_text(calc.cell(i,1),f'{money(b):.2f}' if b!='-' else '-',bold=i in [3,8],size=9,align=WD_ALIGN_PARAGRAPH.RIGHT)

    bottom_cols=2 if not pay_qr else 3
    bottom=doc.add_table(rows=1, cols=bottom_cols); bottom.alignment=WD_TABLE_ALIGNMENT.CENTER
    b0=bottom.rows[0].cells[0]
    set_cell_text(b0, f'Our Bank - {s.get("bank","")}\nAc. No.- {s.get("account","")}\nIFSC Code- {s.get("ifsc","")}\nPlease Draw The Chq. on the Name of “{s.get("name","")}”.', size=9)
    if pay_qr:
        b1=bottom.rows[0].cells[1]
        b1.text=''; b1.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER; b1.paragraphs[0].add_run('Scan to Pay\n').bold=True; b1.paragraphs[0].add_run().add_picture(str(pay_qr), width=Inches(0.95)); set_border(b1)
        sig=bottom.rows[0].cells[2]
    else:
        sig=bottom.rows[0].cells[1]
    set_cell_text(sig, f'For: {s.get("name","")} ,\n\n\nAuthorised Signatory', bold=True, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    p=doc.add_paragraph(); p.add_run(f'Total Invoice Amt. In words – {amount_words_simple(total)} .').bold=True
    doc.add_paragraph('Certified that the particulars given above are true and correct')
    try:
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        test_file = docx_path.parent / '.__btrc_write_test.tmp'
        test_file.write_text('ok', encoding='utf-8')
        test_file.unlink(missing_ok=True)
        doc.save(docx_path)
    except PermissionError:
        raise PermissionError('Invoice could not be saved. Close the existing invoice file if it is open in Word and try again.')
    except OSError as e:
        raise OSError(f'Invoice could not be saved. Check disk space, USB/network drive connection, and save location.\n{e}')

    con=db()
    try:
        con.execute('INSERT INTO invoices(bill_no,date,party,gstin,taxable,discount,freight,cgst,sgst,igst,total,docx_path,created_at) VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)',
            (bill, values['date'], values['party'], values['gstin'], money(values['taxable']), money(values['discount']), money(values['freight']), money(values['cgst']), money(values['sgst']), money(values['igst']), total, str(docx_path), (safe_now() or datetime.now()).isoformat(timespec='seconds')))
        con.commit()
    except sqlite3.IntegrityError:
        raise ValueError(f'Invoice {bill} already exists. Use a new bill number. Overwrite is disabled to protect invoice history.')
    finally:
        con.close()
    return docx_path


def print_file(path):
    """Send generated invoice to printer. Works best on Windows with MS Word installed."""
    path = str(path)
    if platform.system().lower().startswith('win'):
        try:
            os.startfile(path, 'print')
            return True, 'Print command sent to printer.'
        except Exception as e:
            return False, str(e)
    try:
        subprocess.run(['lp', path], check=True)
        return True, 'Print command sent to printer.'
    except Exception as e:
        return False, str(e)

def export_gst_excel():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
    outdir = records_path('GST Exports'); outdir.mkdir(parents=True, exist_ok=True)
    out = outdir / f'GST_REGISTER_{safe_timestamp()}.xlsx'
    con=db(); cur=con.cursor()
    cur.execute('SELECT bill_no,date,party,gstin,taxable,igst,cgst,sgst,total FROM invoices ORDER BY CAST(bill_no AS INTEGER)')
    rows=cur.fetchall(); con.close()
    wb=Workbook(); ws=wb.active; ws.title='GST REGISTER'
    ws.append(['SALES','','GSTIN- '+BUSINESS['gstin'],'DEMO TEXTILE RESEARCH LAB'])
    ws.append(['DATE','BILL NO','NAME','GSTIN','HSN','TAXABLE','IGST','CGST','SGST','TOTAL','RATE'])
    thin=Side(style='thin', color='444444')
    head_fill=PatternFill('solid', fgColor='EDEDED')
    dark_fill=PatternFill('solid', fgColor='1F4E78')
    for c in range(1,12):
        ws.cell(2,c).font=Font(bold=True, color='FFFFFF')
        ws.cell(2,c).fill=dark_fill
        ws.cell(2,c).alignment=Alignment(horizontal='center')
    totals=[0,0,0,0,0]
    for r,row in enumerate(rows, start=3):
        bill,date,party,gstin,taxable,igst,cgst,sgst,total=row
        vals=[date,bill,party,gstin,'998346',taxable,igst,cgst,sgst,total,'18%']
        for c,v in enumerate(vals, start=1): ws.cell(r,c).value=v
        totals[0]+=money(taxable); totals[1]+=money(igst); totals[2]+=money(cgst); totals[3]+=money(sgst); totals[4]+=money(total)
    tr=ws.max_row+1
    ws.cell(tr,5).value='TOTAL'; ws.cell(tr,5).font=Font(bold=True)
    for idx,val in zip([6,7,8,9,10], totals):
        ws.cell(tr,idx).value=val; ws.cell(tr,idx).font=Font(bold=True)
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=11):
        for cell in row:
            cell.border=Border(left=thin,right=thin,top=thin,bottom=thin)
            if isinstance(cell.value,(int,float)): cell.number_format='0.00'
    widths=[12,10,35,20,10,12,10,10,10,12,8]
    for i,w in enumerate(widths, start=1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes='A3'
    wb.save(out)
    return out


def parse_invoice_date(date_text):
    try:
        return datetime.strptime(str(date_text).strip(), '%d/%m/%Y')
    except Exception:
        return None

def get_available_years():
    con=db(); cur=con.cursor(); cur.execute('SELECT date FROM invoices'); rows=cur.fetchall(); con.close()
    years=set()
    for (d,) in rows:
        dt=parse_invoice_date(d)
        if dt: years.add(str(dt.year))
    return sorted(years, reverse=True)

def get_available_months(year):
    con=db(); cur=con.cursor(); cur.execute('SELECT date FROM invoices'); rows=cur.fetchall(); con.close()
    months=[]; seen=set()
    for (d,) in rows:
        dt=parse_invoice_date(d)
        if dt and str(dt.year)==str(year):
            key=dt.strftime('%m'); label=dt.strftime('%B').upper()
            if key not in seen:
                seen.add(key); months.append((key,label))
    return [label for key,label in sorted(months)]

def month_number_from_name(month_name):
    try:
        return datetime.strptime(month_name.title(), '%B').month
    except Exception:
        return None

def invoice_exists(bill_no):
    bill=normalize_bill_no(bill_no)
    con=db(); cur=con.cursor(); cur.execute('SELECT 1 FROM invoices WHERE bill_no=?', (bill,)); ok=cur.fetchone() is not None; con.close(); return ok

def create_backup(label='manual'):
    """Create and verify a FULL backup.

    Includes:
    - SQLite database
    - customer/rate master files
    - generated invoice documents
    - GST exports
    - monthly reports
    - purchase folder records

    Does NOT include Backups folder itself, so backups do not become nested/huge.
    """
    try:
        base = records_base()
        bdir = records_path('Backups'); bdir.mkdir(parents=True, exist_ok=True)
        out = bdir / f'DEMO_BACKUP_{label}_{safe_timestamp()}.zip'
        include_dirs = ['Invoices', 'Purchases', 'GST Exports', 'Monthly Reports']
        with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
            if DB.exists():
                z.write(DB, 'demo_billing.db')
            for f in [MASTER/'customers.xlsx', MASTER/'rate_list.xlsx']:
                if f.exists():
                    z.write(f, f'master/{f.name}')
            for folder in include_dirs:
                src = base / folder
                if src.exists():
                    for file in src.rglob('*'):
                        if file.is_file():
                            z.write(file, str(Path(folder) / file.relative_to(src)))
        if not out.exists() or out.stat().st_size == 0:
            raise OSError('Backup file was not created correctly.')
        with zipfile.ZipFile(out, 'r') as z:
            bad = z.testzip()
            if bad:
                raise OSError(f'Backup verification failed at {bad}.')
            names = set(z.namelist())
            required = {'demo_billing.db', 'master/customers.xlsx', 'master/rate_list.xlsx'}
            missing = [x for x in required if x not in names]
            if missing:
                raise OSError('Backup is incomplete. Missing: ' + ', '.join(missing))
        return out
    except Exception as e:
        print('Backup failed:', e)
        return None

def _safe_extract_zip(z, target_dir):
    """Safely extract a ZIP without allowing path traversal."""
    target_dir = Path(target_dir).resolve()
    for member in z.infolist():
        dest = (target_dir / member.filename).resolve()
        if not str(dest).startswith(str(target_dir)):
            raise ValueError('Unsafe backup ZIP path detected.')
    z.extractall(target_dir)

def restore_backup_from_zip(zip_path):
    """Restore full DEMO backup: DB, masters, invoices, exports, reports, purchases."""
    zip_path = Path(zip_path)
    if not zip_path.exists():
        raise FileNotFoundError('Selected backup file does not exist.')
    if zip_path.suffix.lower() != '.zip':
        raise ValueError('Please select a valid .zip backup file.')

    base = records_base()
    with zipfile.ZipFile(zip_path, 'r') as z:
        bad = z.testzip()
        if bad:
            raise ValueError(f'Backup ZIP is damaged. Problem file: {bad}')
        names = set(z.namelist())
        if 'demo_billing.db' not in names or 'master/customers.xlsx' not in names or 'master/rate_list.xlsx' not in names:
            raise ValueError('This does not look like a complete DEMO backup ZIP.')

    # Safety backup of current system before replacing anything
    safety = create_backup('before_restore_safety')

    temp_dir = DATA / f'_restore_temp_{safe_timestamp()}'
    if temp_dir.exists():
        shutil.rmtree(temp_dir, ignore_errors=True)
    temp_dir.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(zip_path, 'r') as z:
            _safe_extract_zip(z, temp_dir)

        restored = []
        # Restore database and master files
        if (temp_dir/'demo_billing.db').exists():
            shutil.copy2(temp_dir/'demo_billing.db', DB); restored.append('database')
        if (temp_dir/'master'/'customers.xlsx').exists():
            MASTER.mkdir(parents=True, exist_ok=True)
            shutil.copy2(temp_dir/'master'/'customers.xlsx', MASTER/'customers.xlsx'); restored.append('customers')
        if (temp_dir/'master'/'rate_list.xlsx').exists():
            MASTER.mkdir(parents=True, exist_ok=True)
            shutil.copy2(temp_dir/'master'/'rate_list.xlsx', MASTER/'rate_list.xlsx'); restored.append('rate list')

        # Restore records folders. Existing folders are replaced to match backup.
        for folder in ['Invoices', 'Purchases', 'GST Exports', 'Monthly Reports']:
            src = temp_dir / folder
            dst = base / folder
            if src.exists():
                if dst.exists():
                    shutil.rmtree(dst, ignore_errors=True)
                shutil.copytree(src, dst)
                restored.append(folder)

        if not restored:
            raise ValueError('No restorable DEMO files were found in backup.')
        return safety, restored
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

def invoice_rows_filtered(year=None, month=None, search=''):
    con=db(); cur=con.cursor()
    cur.execute('SELECT bill_no,date,party,gstin,taxable,discount,freight,cgst,sgst,igst,total,docx_path FROM invoices')
    all_rows=cur.fetchall(); con.close()
    out=[]; term=(search or '').lower().strip(); month_num=month_number_from_name(month) if month else None
    for row in all_rows:
        bill,date,party,gstin,taxable,discount,freight,cgst,sgst,igst,total,path=row
        dt=parse_invoice_date(date)
        if (year or month_num) and not dt: continue
        if year and str(dt.year)!=str(year): continue
        if month_num and dt.month!=month_num: continue
        if term and term not in str(party).lower() and term not in str(bill).lower(): continue
        out.append(row)
    def bill_key(r):
        try: return int(str(r[0]).lstrip('0') or '0')
        except Exception: return 0
    return sorted(out, key=bill_key)


def save_purchase(values):
    """Save a purchase entry in the local database."""
    dt,msg=validate_invoice_date_str(values.get('date',''))
    if not dt:
        raise ValueError(msg)
    if money(values.get('taxable')) < 0 or money(values.get('cgst')) < 0 or money(values.get('sgst')) < 0 or money(values.get('igst')) < 0 or money(values.get('total')) < 0:
        raise ValueError('Purchase amounts cannot be negative.')
    con=db()
    con.execute("""INSERT INTO purchases(bill_no,date,party,gstin,taxable,cgst,sgst,igst,total,note,created_at)
                   VALUES(?,?,?,?,?,?,?,?,?,?,?)""", (
        str(values.get('bill_no','')).strip(), str(values.get('date','')).strip(), str(values.get('party','')).strip(),
        str(values.get('gstin','')).strip(), money(values.get('taxable')), money(values.get('cgst')),
        money(values.get('sgst')), money(values.get('igst')), money(values.get('total')),
        str(values.get('note','')).strip(), (safe_now() or datetime.now()).isoformat(timespec='seconds')
    ))
    con.commit(); con.close()

def get_available_years_for(kind='sales'):
    table='purchases' if kind=='purchases' else 'invoices'
    con=db(); cur=con.cursor(); cur.execute(f'SELECT date FROM {table}'); rows=cur.fetchall(); con.close()
    years=set()
    for (d,) in rows:
        dt=parse_invoice_date(d)
        if dt: years.add(str(dt.year))
    return sorted(years, reverse=True)

def get_available_months_for(year, kind='sales'):
    if not year: return []
    table='purchases' if kind=='purchases' else 'invoices'
    con=db(); cur=con.cursor(); cur.execute(f'SELECT date FROM {table}'); rows=cur.fetchall(); con.close()
    months=[]; seen=set()
    for (d,) in rows:
        dt=parse_invoice_date(d)
        if dt and str(dt.year)==str(year):
            key=dt.strftime('%m'); label=dt.strftime('%B').upper()
            if key not in seen:
                seen.add(key); months.append((key,label))
    return [label for key,label in sorted(months)]

def purchase_rows_filtered(year=None, month=None, search=''):
    con=db(); cur=con.cursor()
    cur.execute('SELECT bill_no,date,party,gstin,taxable,cgst,sgst,igst,total,note FROM purchases')
    all_rows=cur.fetchall(); con.close()
    out=[]; term=(search or '').lower().strip(); month_num=month_number_from_name(month) if month else None
    for row in all_rows:
        bill,date,party,gstin,taxable,cgst,sgst,igst,total,note=row
        dt=parse_invoice_date(date)
        if (year or month_num) and not dt: continue
        if year and str(dt.year)!=str(year): continue
        if month_num and dt.month!=month_num: continue
        if term and term not in str(party).lower() and term not in str(bill).lower(): continue
        out.append(row)
    def bill_key(r):
        try: return int(str(r[0]).lstrip('0') or '0')
        except Exception: return 0
    return sorted(out, key=bill_key)

def export_purchase_excel_filtered(year=None, month=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
    outdir=records_path('GST Exports'); outdir.mkdir(parents=True, exist_ok=True)
    suffix=((str(year) if year else 'ALL') + ('_'+str(month).replace(' ','_') if month else ''))
    out=outdir/f'PURCHASE_REGISTER_{suffix}_{safe_timestamp()}.xlsx'
    rows=purchase_rows_filtered(year, month)
    wb=Workbook(); ws=wb.active; ws.title='PURCHASE REGISTER'
    ws.append(['PURCHASES','','GSTIN- '+BUSINESS['gstin'],'DEMO TEXTILE RESEARCH LAB'])
    ws.append(['YEAR', year or 'ALL', 'MONTH', month or 'ALL'])
    ws.append(['DATE','BILL NO','NAME','GSTIN','TAXABLE','IGST','CGST','SGST','TOTAL','NOTE'])
    thin=Side(style='thin', color='444444'); dark_fill=PatternFill('solid', fgColor='1F4E78')
    for c in range(1,11):
        ws.cell(3,c).font=Font(bold=True, color='FFFFFF'); ws.cell(3,c).fill=dark_fill; ws.cell(3,c).alignment=Alignment(horizontal='center')
    totals=[0,0,0,0,0]
    for r,row in enumerate(rows, start=4):
        bill,date,party,gstin,taxable,cgst,sgst,igst,total,note=row
        vals=[date,bill,party,gstin,taxable,igst,cgst,sgst,total,note]
        for c,v in enumerate(vals, start=1): ws.cell(r,c).value=v
        totals[0]+=money(taxable); totals[1]+=money(igst); totals[2]+=money(cgst); totals[3]+=money(sgst); totals[4]+=money(total)
    tr=ws.max_row+1; ws.cell(tr,4).value='TOTAL'; ws.cell(tr,4).font=Font(bold=True)
    for idx,val in zip([5,6,7,8,9], totals): ws.cell(tr,idx).value=val; ws.cell(tr,idx).font=Font(bold=True)
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=10):
        for cell in row:
            cell.border=Border(left=thin,right=thin,top=thin,bottom=thin)
            if isinstance(cell.value,(int,float)): cell.number_format='0.00'
    widths=[12,12,35,20,12,10,10,10,12,30]
    for i,w in enumerate(widths, start=1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes='A4'; wb.save(out); return out

def export_gst_excel_filtered(year=None, month=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
    outdir=records_path('GST Exports'); outdir.mkdir(parents=True, exist_ok=True)
    suffix=((str(year) if year else 'ALL') + ('_'+str(month).replace(' ','_') if month else ''))
    out=outdir/f'GST_REGISTER_{suffix}_{safe_timestamp()}.xlsx'
    rows=invoice_rows_filtered(year, month)
    wb=Workbook(); ws=wb.active; ws.title='GST REGISTER'
    ws.append(['SALES','','GSTIN- '+BUSINESS['gstin'],'DEMO TEXTILE RESEARCH LAB'])
    ws.append(['YEAR', year or 'ALL', 'MONTH', month or 'ALL'])
    ws.append(['DATE','BILL NO','NAME','GSTIN','HSN','TAXABLE','IGST','CGST','SGST','TOTAL','RATE'])
    thin=Side(style='thin', color='444444'); dark_fill=PatternFill('solid', fgColor='1F4E78')
    for c in range(1,12):
        ws.cell(3,c).font=Font(bold=True, color='FFFFFF'); ws.cell(3,c).fill=dark_fill; ws.cell(3,c).alignment=Alignment(horizontal='center')
    totals=[0,0,0,0,0]
    for r,row in enumerate(rows, start=4):
        bill,date,party,gstin,taxable,discount,freight,cgst,sgst,igst,total,path=row
        tax=money(igst)+money(cgst)+money(sgst); rate=(tax/money(taxable)*100) if money(taxable) else 0
        vals=[date,bill,party,gstin,'998346',taxable,igst,cgst,sgst,total,(f'{rate:.0f}%' if abs(rate-round(rate))<0.05 else f'{rate:.2f}%')]
        for c,v in enumerate(vals, start=1): ws.cell(r,c).value=v
        totals[0]+=money(taxable); totals[1]+=money(igst); totals[2]+=money(cgst); totals[3]+=money(sgst); totals[4]+=money(total)
    tr=ws.max_row+1; ws.cell(tr,5).value='TOTAL'; ws.cell(tr,5).font=Font(bold=True)
    for idx,val in zip([6,7,8,9,10], totals): ws.cell(tr,idx).value=val; ws.cell(tr,idx).font=Font(bold=True)
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=11):
        for cell in row:
            cell.border=Border(left=thin,right=thin,top=thin,bottom=thin)
            if isinstance(cell.value,(int,float)): cell.number_format='0.00'
    widths=[12,10,35,20,10,12,10,10,10,12,8]
    for i,w in enumerate(widths, start=1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes='A4'; wb.save(out); return out

def export_report_filtered(year=None, month=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
    sales=invoice_rows_filtered(year, month); purchases=purchase_rows_filtered(year, month)
    outdir=records_path('Monthly Reports'); outdir.mkdir(parents=True, exist_ok=True)
    suffix=((str(year) if year else 'ALL') + ('_'+str(month).replace(' ','_') if month else ''))
    out=outdir/f'MONTHLY_SUMMARY_{suffix}_{safe_timestamp()}.xlsx'
    sales_taxable=sum(money(r[4]) for r in sales); sales_total=sum(money(r[10]) for r in sales); out_gst=sum(money(r[7])+money(r[8])+money(r[9]) for r in sales)
    purchase_taxable=sum(money(r[4]) for r in purchases); purchase_total=sum(money(r[8]) for r in purchases); in_gst=sum(money(r[5])+money(r[6])+money(r[7]) for r in purchases)
    profit=sales_taxable-purchase_taxable
    wb=Workbook(); ws=wb.active; ws.title='SUMMARY'
    ws.append(['DEMO MONTHLY REPORT'])
    ws.append(['Year', year or 'ALL', 'Month', month or 'ALL'])
    ws.append([])
    data=[('Total Sales Bills',len(sales)),('Total Sales Amount',sales_total),('GST Collected from Customers',out_gst),('', ''),('Total Purchase Records',len(purchases)),('Total Purchase Amount',purchase_total),('GST Paid on Purchases',in_gst),('', ''),('Final GST To Pay',out_gst-in_gst),('Estimated Earnings',profit)]
    ws.append(['Metric','Value'])
    for a,b in data: ws.append([a,b])
    thin=Side(style='thin', color='444444'); dark_fill=PatternFill('solid', fgColor='1F4E78')
    ws['A1'].font=Font(bold=True,size=16); ws['A4'].font=Font(bold=True,color='FFFFFF'); ws['B4'].font=Font(bold=True,color='FFFFFF'); ws['A4'].fill=dark_fill; ws['B4'].fill=dark_fill
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=2):
        for cell in row:
            cell.border=Border(left=thin,right=thin,top=thin,bottom=thin)
            if isinstance(cell.value,(int,float)): cell.number_format='0.00'
    ws.column_dimensions['A'].width=28; ws.column_dimensions['B'].width=18
    wb.save(out); return out


def validate_invoice_values(vals):
    try:
        vals['bill_no'] = normalize_bill_no(vals.get('bill_no',''))
    except ValueError as e:
        messagebox.showerror('Invalid Bill No.', str(e))
        return False
    dt,msg=validate_invoice_date_str(vals.get('date',''))
    if not dt:
        messagebox.showerror('Invalid Date', msg)
        return False
    today=safe_now()
    if today and dt.date() > today.date():
        if not messagebox.askyesno('Future Date', 'Invoice date is in the future. Continue only if this is intentional?'):
            return False
    if not str(vals.get('address','')).strip():
        messagebox.showerror('Missing Address','Customer address is required for invoice.')
        return False
    gstin=str(vals.get('gstin','')).strip().upper()
    if not gstin or not validate_gstin(gstin):
        messagebox.showerror('Invalid GSTIN','Customer GSTIN is required and must be valid.')
        return False
    code=str(vals.get('code','')).strip().zfill(2)
    if code not in VALID_STATE_CODES:
        messagebox.showerror('Invalid State Code','State code is invalid. Example: Rajasthan = 08.')
        return False
    gst_state=state_code_from_gstin(gstin)
    if gst_state and gst_state != code:
        messagebox.showerror('GSTIN/State Mismatch', f'GSTIN starts with state code {gst_state}, but entered state code is {code}. Please correct it.')
        return False
    gross=money(vals.get('gross'))
    discount=money(vals.get('discount'))
    if discount < 0:
        messagebox.showerror('Invalid Discount','Discount cannot be negative.')
        return False
    if discount > gross:
        messagebox.showerror('Invalid Discount', 'Discount cannot be greater than the total item amount.')
        return False
    if money(vals.get('freight')) < 0:
        messagebox.showerror('Invalid Freight','Freight cannot be negative.')
        return False
    seen_tests=set()
    for idx,it in enumerate(vals.get('items',[]), start=1):
        key=(str(it.get('test','')).strip().lower(), str(it.get('operation','')).strip().lower())
        if key in seen_tests:
            if not messagebox.askyesno('Duplicate Test', f'Test row {idx} looks duplicate. Continue anyway?'):
                return False
        seen_tests.add(key)
        if not str(it.get('hsn','')).strip():
            messagebox.showerror('Missing HSN', f'HSN code is required in test row {idx}.')
            return False
        if money(it.get('qty')) <= 0:
            messagebox.showerror('Invalid Quantity', f'Quantity must be greater than 0 in test row {idx}.')
            return False
        if money(it.get('amount')) < 0:
            messagebox.showerror('Invalid Amount', f'Amount cannot be negative in test row {idx}.')
            return False
    return True

class App(tk.Tk):
    def __init__(self):
        super().__init__(); self.title(APP); self.geometry('1280x720'); self.configure(bg='#0f172a'); db(); self.login()
    def clear(self):
        for w in self.winfo_children(): w.destroy()
    def login(self):
        self.clear(); frm=tk.Frame(self,bg='#0f172a'); frm.pack(expand=True)
        tk.Label(frm,text='GST Billing Demo',font=('Segoe UI',28,'bold'),fg='white',bg='#0f172a').pack(pady=15)
        tk.Label(frm,text='Enter Password',font=('Segoe UI',11),fg='#cbd5e1',bg='#0f172a').pack()
        e=ttk.Entry(frm,show='*',width=35); e.pack(pady=10); e.focus()
        self.login_attempts=getattr(self,'login_attempts',0)
        def go(event=None):
            s=get_settings(); stored=s.get('password_hash') or s.get('password', DEFAULT_PASSWORD)
            if self.login_attempts >= 5:
                messagebox.showerror('Locked','Too many wrong attempts. Close and reopen the app after a few minutes.')
                return
            if check_password(e.get(), stored):
                self.login_attempts=0; self.home()
            else:
                self.login_attempts+=1; messagebox.showerror('Wrong password',f'Incorrect password. Attempts left: {max(0,5-self.login_attempts)}')
        e.bind('<Return>',go); ttk.Button(frm,text='Open App',command=go).pack(pady=8)
    def home(self):
        self.clear(); side=tk.Frame(self,bg='#111827',width=210); side.pack(side='left',fill='y'); main=tk.Frame(self,bg='#f8fafc'); main.pack(side='right',fill='both',expand=True)
        tk.Label(side,text='DEMO',font=('Segoe UI',20,'bold'),fg='white',bg='#111827').pack(pady=25)
        for txt,cmd in [('Create Invoice',lambda:self.invoice(main)),('Invoices',lambda:self.invoices(main)),('Purchases',lambda:self.purchases(main)),('GST Export',lambda:self.gst_export(main)),('Monthly Reports',lambda:self.reports(main)),('Customers File',self.open_customers),('Rate List File',self.open_rates),('Settings',lambda:self.settings(main))]:
            tk.Button(side,text=txt,command=cmd,bg='#1f2937',fg='white',relief='flat',font=('Segoe UI',11),pady=10).pack(fill='x',padx=12,pady=5)
        self.invoice(main)
    def wipe(self,main):
        for w in main.winfo_children(): w.destroy()
    def entries_return_next(self, widgets):
        for i,w in enumerate(widgets[:-1]): w.bind('<Return>', lambda e, nxt=widgets[i+1]: (nxt.focus_set(), 'break'))
    def invoice(self,main):
        self.wipe(main)
        tk.Label(main,text='Create Invoice',font=('Segoe UI',21,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=(10,6))
        box=tk.Frame(main,bg='white'); box.pack(fill='both',expand=True,padx=25,pady=(0,8))
        self.customers=load_customers(); self.rates=load_rates(); self.items=[]
        fields={}
        labels=['Bill No','Date (dd/mm/yyyy)','Party Name','Address','GSTIN','State','Code','Phone','Discount','Freight']
        grid=tk.Frame(box,bg='white'); grid.pack(padx=15,pady=(6,4),fill='x')
        for i,lbl in enumerate(labels):
            tk.Label(grid,text=lbl,bg='white',anchor='w').grid(row=i//2,column=(i%2)*2,sticky='w',padx=8,pady=3)
            ent=ttk.Combobox(grid,width=42) if lbl=='Party Name' else ttk.Entry(grid,width=45)
            ent.grid(row=i//2,column=(i%2)*2+1,padx=8,pady=3); fields[lbl]=ent
        fields['Date (dd/mm/yyyy)'].insert(0, safe_date_default())
        fields['State'].insert(0,'Rajasthan'); fields['Code'].insert(0,'08'); fields['Discount'].insert(0,'0'); fields['Freight'].insert(0,'0')
        fields['Party Name']['values']=[c.get('name','') for c in self.customers]

        def refresh_customers():
            self.customers=load_customers(); fields['Party Name']['values']=[c.get('name','') for c in self.customers]
        def filter_customers(event=None):
            text=fields['Party Name'].get().lower()
            vals=[str(c.get('name','')) for c in self.customers if text in str(c.get('name','')).lower()]
            fields['Party Name']['values']=vals[:30]
        def fill_customer(event=None):
            name=fields['Party Name'].get().strip().lower()
            for c in self.customers:
                if str(c.get('name','')).strip().lower()==name:
                    mapping=[('Address','address'),('GSTIN','gstin'),('State','state'),('Code','state code'),('Phone','phone')]
                    for key, col in mapping:
                        fields[key].delete(0,'end'); fields[key].insert(0,str(c.get(col,'') or ('08' if key=='Code' else '')))
                    break
        fields['Party Name'].bind('<KeyRelease>',lambda e:(filter_customers(), fields['Party Name'].event_generate('<Down>'))); fields['Party Name'].bind('<<ComboboxSelected>>',fill_customer); fields['Party Name'].bind('<Return>',lambda e:(fill_customer(), fields['Address'].focus_set(), 'break'))

        def add_customer_popup():
            win=tk.Toplevel(self); win.title('Add Customer'); win.geometry('480x360'); ents={}
            for i,k in enumerate(['name','address','gstin','state','state code','phone']):
                tk.Label(win,text=k.upper()).grid(row=i,column=0,sticky='w',padx=10,pady=6)
                e=ttk.Entry(win,width=45); e.grid(row=i,column=1,padx=10,pady=6); ents[k]=e
            ents['state'].insert(0,'Rajasthan'); ents['state code'].insert(0,'08')
            def save():
                if not ents['name'].get().strip(): messagebox.showwarning('Missing','Customer name required'); return
                try:
                    add_customer_to_file({k:e.get() for k,e in ents.items()}); create_backup('after_customer')
                    refresh_customers(); fields['Party Name'].delete(0,'end'); fields['Party Name'].insert(0,ents['name'].get()); fill_customer(); win.destroy()
                except Exception as e:
                    messagebox.showerror('Customer Error', str(e))
            ttk.Button(win,text='Save Customer',command=save).grid(row=7,column=1,sticky='e',pady=15)
        ttk.Button(grid,text='Add Customer in Customer List',command=add_customer_popup).grid(row=5,column=3,sticky='e',padx=8,pady=3)

        itembox=tk.LabelFrame(box,text='Tests',bg='white',font=('Segoe UI',11,'bold')); itembox.pack(fill='x',padx=18,pady=(4,4))
        test_cb=ttk.Combobox(itembox,width=38); op_e=ttk.Entry(itembox,width=32); hsn_e=ttk.Entry(itembox,width=12); qty_e=ttk.Entry(itembox,width=8); amt_e=ttk.Entry(itembox,width=12)
        for i,(lab,w) in enumerate([('Test',test_cb),('Report No.',op_e),('HSN',hsn_e),('Qty',qty_e),('Amount',amt_e)]):
            tk.Label(itembox,text=lab,bg='white').grid(row=0,column=i,padx=6,pady=2,sticky='w'); w.grid(row=1,column=i,padx=6,pady=2)
        qty_e.insert(0,'1')
        test_cb['values']=[r.get('test','') or r.get('test parameter','') for r in self.rates]
        def filter_tests(event=None):
            text=test_cb.get().lower()
            vals=[str(r.get('test','') or r.get('test parameter','')) for r in self.rates if text in str(r.get('test','') or r.get('test parameter','')).lower()]
            test_cb['values']=vals[:40]
        def fill_rate(event=None):
            p=test_cb.get().strip().lower()
            for r in self.rates:
                nm=str(r.get('test','') or r.get('test parameter','')).strip().lower()
                if nm==p:
                    for e,val in [(op_e,r.get('operation/report','') or ''),(hsn_e,r.get('hsn code','') or ''),(qty_e,r.get('quantity','1') or '1'),(amt_e,r.get('amount','') or '')]:
                        e.delete(0,'end'); e.insert(0,str(val))
                    break
        test_cb.bind('<KeyRelease>',lambda e:(filter_tests(), test_cb.event_generate('<Down>'))); test_cb.bind('<<ComboboxSelected>>',fill_rate); test_cb.bind('<Return>',lambda e:(fill_rate(),op_e.focus_set(),'break'))
        self.entries_return_next([op_e, hsn_e, qty_e, amt_e])

        tree=ttk.Treeview(box,columns=('sn','test','operation','hsn','qty','amount','total'),show='headings',height=5)
        heading_names={'sn':'SN','test':'TEST','operation':'REPORT NO.','hsn':'HSN','qty':'QTY','amount':'AMOUNT','total':'TOTAL'}
        for c,wid in [('sn',45),('test',280),('operation',240),('hsn',90),('qty',70),('amount',110),('total',110)]:
            tree.heading(c,text=heading_names[c]); tree.column(c,width=wid, anchor='center' if c in ['sn','hsn','qty','amount','total'] else 'w')
        tree.pack(fill='x',padx=18,pady=(4,3))
        total_lbl=tk.Label(box,text='Taxable: ₹0 | CGST: ₹0 | SGST: ₹0 | IGST: ₹0 | Total: ₹0',font=('Segoe UI',13,'bold'),bg='white',fg='#1d4ed8'); total_lbl.pack(anchor='w',padx=22,pady=(4,4))
        calc_vals={}
        def refresh_items():
            for r in tree.get_children(): tree.delete(r)
            for i,it in enumerate(self.items,1):
                tree.insert('', 'end', values=(i,it['test'],it['operation'],it['hsn'],it['qty'],f'{it["amount"]:.2f}',f'{it["line_total"]:.2f}'))
            calc()
        def add_item(event=None):
            if not test_cb.get().strip() or not amt_e.get().strip():
                messagebox.showwarning('Missing','Test and amount are required'); return 'break'
            if not hsn_e.get().strip():
                messagebox.showwarning('Missing HSN','HSN code is required for this test.'); return 'break'
            qty=money(qty_e.get() or 1); amount=money(amt_e.get())
            if qty <= 0:
                messagebox.showwarning('Invalid Qty','Quantity must be greater than 0.'); return 'break'
            if amount < 0:
                messagebox.showwarning('Invalid Amount','Amount cannot be negative.'); return 'break'
            line_total=round(qty*amount,2)
            self.items.append({'test':test_cb.get().strip(),'operation':op_e.get().strip(),'hsn':hsn_e.get().strip(),'qty':qty_e.get().strip() or '1','amount':amount,'line_total':line_total})
            for e in [test_cb,op_e,hsn_e,qty_e,amt_e]: e.delete(0,'end')
            qty_e.insert(0,'1'); test_cb.focus_set(); refresh_items(); return 'break'
        def remove_item():
            sel=tree.focus()
            if sel:
                idx=tree.index(sel); self.items.pop(idx); refresh_items()
        btns=tk.Frame(itembox,bg='white'); btns.grid(row=1,column=5,padx=8)
        ttk.Button(btns,text='Add',command=add_item).pack(fill='x',pady=2)
        ttk.Button(btns,text='Remove Selected',command=remove_item).pack(fill='x',pady=2)
        amt_e.bind('<Return>', add_item)
        def calc(event=None):
            gross=sum(money(i.get('line_total')) for i in self.items)
            disc=money(fields['Discount'].get()); freight=money(fields['Freight'].get())
            taxable=gross - disc + freight
            code=fields['Code'].get().strip() or '08'
            if code=='08': cgst=sgst=round(taxable*0.09,2); igst=0
            else: igst=round(taxable*0.18,2); cgst=sgst=0
            total=round(taxable+cgst+sgst+igst,2)
            calc_vals.update(gross=gross,taxable=taxable,cgst=cgst,sgst=sgst,igst=igst,total=total)
            total_lbl.config(text=f'Amount: ₹{gross:.2f} | Taxable: ₹{taxable:.2f} | CGST: ₹{cgst:.2f} | SGST: ₹{sgst:.2f} | IGST: ₹{igst:.2f} | Total: ₹{total:.2f}')
        for k in ['Discount','Freight','Code']: fields[k].bind('<KeyRelease>',calc)
        self.entries_return_next([fields['Bill No'],fields['Date (dd/mm/yyyy)'],fields['Party Name'],fields['Address'],fields['GSTIN'],fields['State'],fields['Code'],fields['Phone'],fields['Discount'],fields['Freight'],test_cb])
        def generate():
            calc(); fill_customer()
            vals={
                'bill_no':fields['Bill No'].get(), 'date':fields['Date (dd/mm/yyyy)'].get(), 'party':fields['Party Name'].get(), 'address':fields['Address'].get(), 'gstin':fields['GSTIN'].get(),
                'state':fields['State'].get(), 'code':fields['Code'].get(), 'phone':fields['Phone'].get(), 'items':self.items[:],
                'discount':fields['Discount'].get(), 'freight':fields['Freight'].get(), **calc_vals
            }
            if not vals['bill_no'] or not vals['party'] or not vals['items']:
                messagebox.showwarning('Missing','Bill no, party name and at least one test are required.'); return
            if not validate_invoice_values(vals): return
            if invoice_exists(vals['bill_no']):
                messagebox.showerror('Duplicate Bill No.', f'Invoice {normalize_bill_no(vals["bill_no"])} already exists. Please use a new bill number. Overwrite is disabled to protect invoice history.')
                return
            try:
                path=invoice_docx(vals); create_backup('after_invoice'); self.last_invoice_path=path; messagebox.showinfo('Done',f'Invoice saved:\n{path}'); webbrowser.open(str(path))
            except PermissionError as e: messagebox.showerror('File Open',str(e))
            except OSError as e: messagebox.showerror('Save Error',str(e))
            except Exception as e: messagebox.showerror('Error',str(e))
        action_bar=tk.Frame(box,bg='white'); action_bar.pack(fill='x',padx=22,pady=(2,6))
        right_actions=tk.Frame(action_bar,bg='white'); right_actions.pack(side='right')
        self.last_invoice_path=None
        def generate_and_remember():
            generate()
        def generate_print():
            calc(); fill_customer()
            vals={
                'bill_no':fields['Bill No'].get(), 'date':fields['Date (dd/mm/yyyy)'].get(), 'party':fields['Party Name'].get(), 'address':fields['Address'].get(), 'gstin':fields['GSTIN'].get(),
                'state':fields['State'].get(), 'code':fields['Code'].get(), 'phone':fields['Phone'].get(), 'items':self.items[:],
                'discount':fields['Discount'].get(), 'freight':fields['Freight'].get(), **calc_vals
            }
            if not vals['bill_no'] or not vals['party'] or not vals['items']:
                messagebox.showwarning('Missing','Bill no, party name and at least one test are required.'); return
            if not validate_invoice_values(vals): return
            if invoice_exists(vals['bill_no']):
                messagebox.showerror('Duplicate Bill No.', f'Invoice {normalize_bill_no(vals["bill_no"])} already exists. Please use a new bill number. Overwrite is disabled to protect invoice history.')
                return
            try:
                path=invoice_docx(vals); create_backup('after_invoice'); self.last_invoice_path=path
                ok,msg=print_file(path)
                if ok: messagebox.showinfo('Print',f'Invoice saved and print command sent:\n{path}')
                else: messagebox.showerror('Print Error',msg + '\n\nCheck that Microsoft Word is installed and your default printer is correct.')
            except PermissionError as e: messagebox.showerror('File Open',str(e))
            except OSError as e: messagebox.showerror('Save Error',str(e))
            except Exception as e: messagebox.showerror('Error',str(e))
        ttk.Button(right_actions,text='Generate Invoice',command=generate,width=18).pack(side='left',padx=8)
        ttk.Button(right_actions,text='Generate & Print',command=generate_print,width=18).pack(side='left',padx=8)
    def invoices(self,main):
        self.wipe(main)
        tk.Label(main,text='Invoice Records',font=('Segoe UI',22,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=15)
        top=tk.Frame(main,bg='#f8fafc'); top.pack(fill='x',padx=25,pady=5)
        tk.Label(top,text='Year:',bg='#f8fafc').pack(side='left',padx=(0,5))
        year_cb=ttk.Combobox(top,width=12,state='readonly'); year_cb.pack(side='left',padx=5)
        tk.Label(top,text='Month:',bg='#f8fafc').pack(side='left',padx=(15,5))
        month_cb=ttk.Combobox(top,width=18,state='readonly'); month_cb.pack(side='left',padx=5)
        tk.Label(top,text='Search:',bg='#f8fafc').pack(side='left',padx=(18,5))
        q=ttk.Entry(top,width=35); q.pack(side='left',padx=5)
        tree=ttk.Treeview(main,columns=('bill','date','party','gstin','total','path'),show='headings')
        for c,w in [('bill',90),('date',110),('party',300),('gstin',170),('total',110),('path',360)]:
            tree.heading(c,text=c.upper()); tree.column(c,width=w)
        tree.pack(fill='both',expand=True,padx=25,pady=15)
        years=get_available_years(); year_cb['values']=years
        if years: year_cb.set(years[0]); month_cb['values']=get_available_months(years[0])
        if month_cb['values']: month_cb.set(month_cb['values'][0])
        def refresh(event=None):
            for r in tree.get_children(): tree.delete(r)
            year=year_cb.get() or None; month=month_cb.get() or None
            for row in invoice_rows_filtered(year, month, q.get()):
                bill,date,party,gstin,taxable,discount,freight,cgst,sgst,igst,total,path=row
                tree.insert('', 'end', values=(bill,date,party,gstin,f'{money(total):.2f}',path))
        def year_changed(event=None):
            months=get_available_months(year_cb.get()); month_cb['values']=months
            month_cb.set(months[0] if months else '')
            refresh()
        def open_sel(event=None):
            sel=tree.focus()
            if sel: webbrowser.open(tree.item(sel)['values'][5])
        year_cb.bind('<<ComboboxSelected>>',year_changed); month_cb.bind('<<ComboboxSelected>>',refresh); q.bind('<KeyRelease>',refresh); tree.bind('<Double-1>',open_sel)
        refresh()

    def purchases(self,main):
        self.wipe(main)
        tk.Label(main,text='Purchase Records',font=('Segoe UI',22,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=15)
        form=tk.Frame(main,bg='white'); form.pack(fill='x',padx=25,pady=8)
        fields={}
        labels=['Bill No','Date (dd/mm/yyyy)','Party Name','GSTIN','Taxable','CGST','SGST','IGST','Total','Note']
        for i,lbl in enumerate(labels):
            tk.Label(form,text=lbl,bg='white').grid(row=i//5,column=(i%5)*2,sticky='w',padx=8,pady=3)
            e=ttk.Entry(form,width=22); e.grid(row=i//5,column=(i%5)*2+1,padx=5,pady=5); fields[lbl]=e
        fields['Date (dd/mm/yyyy)'].insert(0, safe_date_default())
        def calc_total(event=None):
            if fields['Total'].get().strip(): return
            total=money(fields['Taxable'].get())+money(fields['CGST'].get())+money(fields['SGST'].get())+money(fields['IGST'].get())
            fields['Total'].delete(0,'end'); fields['Total'].insert(0,f'{total:.2f}')
        for k in ['Taxable','CGST','SGST','IGST']: fields[k].bind('<FocusOut>',calc_total)
        self.entries_return_next([fields[k] for k in labels])
        def save_rec():
            if not fields['Bill No'].get().strip() or not fields['Party Name'].get().strip():
                messagebox.showwarning('Missing','Bill no and party name are required'); return
            calc_total()
            vals={'bill_no':fields['Bill No'].get(),'date':fields['Date (dd/mm/yyyy)'].get(),'party':fields['Party Name'].get(),'gstin':fields['GSTIN'].get(),'taxable':fields['Taxable'].get(),'cgst':fields['CGST'].get(),'sgst':fields['SGST'].get(),'igst':fields['IGST'].get(),'total':fields['Total'].get(),'note':fields['Note'].get()}
            if vals['gstin'] and not validate_gstin(vals['gstin']):
                messagebox.showerror('Invalid GSTIN','Purchase GSTIN format is invalid.'); return
            try:
                save_purchase(vals); create_backup('after_purchase'); messagebox.showinfo('Saved','Purchase record saved'); refresh()
                for e in fields.values(): e.delete(0,'end')
                fields['Date (dd/mm/yyyy)'].insert(0, safe_date_default())
            except Exception as e: messagebox.showerror('Error',str(e))
        ttk.Button(form,text='Save Purchase Record',command=save_rec).grid(row=3,column=0,columnspan=2,sticky='w',padx=8,pady=12)
        filt=tk.Frame(main,bg='#f8fafc'); filt.pack(fill='x',padx=25,pady=5)
        tk.Label(filt,text='Year:',bg='#f8fafc').pack(side='left'); year_cb=ttk.Combobox(filt,width=12,state='readonly'); year_cb.pack(side='left',padx=5)
        tk.Label(filt,text='Month:',bg='#f8fafc').pack(side='left',padx=(15,0)); month_cb=ttk.Combobox(filt,width=18,state='readonly'); month_cb.pack(side='left',padx=5)
        tk.Label(filt,text='Search:',bg='#f8fafc').pack(side='left',padx=(15,0)); q=ttk.Entry(filt,width=35); q.pack(side='left',padx=5)
        tree=ttk.Treeview(main,columns=('bill','date','party','gstin','taxable','cgst','sgst','igst','total'),show='headings')
        for c,w in [('bill',80),('date',100),('party',260),('gstin',160),('taxable',100),('cgst',80),('sgst',80),('igst',80),('total',100)]:
            tree.heading(c,text=c.upper()); tree.column(c,width=w)
        tree.pack(fill='both',expand=True,padx=25,pady=12)
        years=get_available_years_for('purchases'); year_cb['values']=years
        if years: year_cb.set(years[0]); month_cb['values']=get_available_months_for(years[0],'purchases')
        if month_cb['values']: month_cb.set(month_cb['values'][0])
        def refresh(event=None):
            for r in tree.get_children(): tree.delete(r)
            for row in purchase_rows_filtered(year_cb.get() or None, month_cb.get() or None, q.get()):
                bill,date,party,gstin,taxable,cgst,sgst,igst,total,note=row
                tree.insert('', 'end', values=(bill,date,party,gstin,f'{money(taxable):.2f}',f'{money(cgst):.2f}',f'{money(sgst):.2f}',f'{money(igst):.2f}',f'{money(total):.2f}'))
        def year_changed(event=None):
            months=get_available_months_for(year_cb.get(),'purchases'); month_cb['values']=months; month_cb.set(months[0] if months else ''); refresh()
        year_cb.bind('<<ComboboxSelected>>',year_changed); month_cb.bind('<<ComboboxSelected>>',refresh); q.bind('<KeyRelease>',refresh); refresh()

    def gst_export(self,main):
        self.wipe(main)
        tk.Label(main,text='GST Export',font=('Segoe UI',22,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=15)
        box=tk.Frame(main,bg='white'); box.pack(fill='x',padx=25,pady=10)
        tk.Label(box,text='Select record type, year and month, then export GST register for that period.',bg='white',font=('Segoe UI',12)).grid(row=0,column=0,columnspan=6,sticky='w',padx=18,pady=12)
        tk.Label(box,text='Type',bg='white').grid(row=1,column=0,sticky='w',padx=18,pady=8)
        type_cb=ttk.Combobox(box,width=15,state='readonly',values=['Sales','Purchases']); type_cb.grid(row=1,column=1,sticky='w',padx=5,pady=8); type_cb.set('Sales')
        tk.Label(box,text='Year',bg='white').grid(row=1,column=2,sticky='w',padx=18,pady=8)
        year_cb=ttk.Combobox(box,width=15,state='readonly'); year_cb.grid(row=1,column=3,sticky='w',padx=5,pady=8)
        tk.Label(box,text='Month',bg='white').grid(row=1,column=4,sticky='w',padx=18,pady=8)
        month_cb=ttk.Combobox(box,width=20,state='readonly'); month_cb.grid(row=1,column=5,sticky='w',padx=5,pady=8)
        count_lbl=tk.Label(box,text='',bg='white',font=('Segoe UI',11,'bold'),fg='#1d4ed8'); count_lbl.grid(row=2,column=0,columnspan=6,sticky='w',padx=18,pady=8)
        def current_type(): return 'sales' if type_cb.get()=='Sales' else 'purchases'
        def reload_periods(event=None):
            years=get_available_years_for(current_type()); year_cb['values']=years
            year_cb.set(years[0] if years else '')
            months=get_available_months_for(year_cb.get(), current_type()) if years else []
            month_cb['values']=months; month_cb.set(months[0] if months else '')
            update_count()
        def update_count(event=None):
            if current_type()=='sales': rows=invoice_rows_filtered(year_cb.get() or None, month_cb.get() or None)
            else: rows=purchase_rows_filtered(year_cb.get() or None, month_cb.get() or None)
            count_lbl.config(text=f'{type_cb.get()} records selected: {len(rows)}')
        def year_changed(event=None):
            months=get_available_months_for(year_cb.get(), current_type()); month_cb['values']=months; month_cb.set(months[0] if months else ''); update_count()
        def do_export():
            try:
                if current_type()=='sales': path=export_gst_excel_filtered(year_cb.get() or None, month_cb.get() or None)
                else: path=export_purchase_excel_filtered(year_cb.get() or None, month_cb.get() or None)
                messagebox.showinfo('Exported',f'{type_cb.get()} GST Excel saved:\n{path}')
                webbrowser.open(str(path))
            except Exception as e: messagebox.showerror('Export Error',str(e))
        type_cb.bind('<<ComboboxSelected>>',reload_periods); year_cb.bind('<<ComboboxSelected>>',year_changed); month_cb.bind('<<ComboboxSelected>>',update_count)
        reload_periods()
        ttk.Button(box,text='Export Selected Month Excel',command=do_export).grid(row=3,column=0,columnspan=2,sticky='w',padx=18,pady=15)

    def reports(self,main):
        self.wipe(main)
        tk.Label(main,text='Monthly Reports',font=('Segoe UI',22,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=15)
        box=tk.Frame(main,bg='white'); box.pack(fill='x',padx=25,pady=10)
        tk.Label(box,text='Simple monthly summary: Sales, Purchases, GST to Pay and Estimated Earnings.',bg='white',font=('Segoe UI',12)).grid(row=0,column=0,columnspan=4,sticky='w',padx=18,pady=12)
        tk.Label(box,text='Year',bg='white').grid(row=1,column=0,sticky='w',padx=18,pady=8)
        year_cb=ttk.Combobox(box,width=15,state='readonly'); year_cb.grid(row=1,column=1,sticky='w',padx=5,pady=8)
        tk.Label(box,text='Month',bg='white').grid(row=1,column=2,sticky='w',padx=18,pady=8)
        month_cb=ttk.Combobox(box,width=20,state='readonly'); month_cb.grid(row=1,column=3,sticky='w',padx=5,pady=8)
        summary=tk.Label(box,text='',bg='white',font=('Consolas',12,'bold'),fg='#0f172a',justify='left'); summary.grid(row=2,column=0,columnspan=4,sticky='w',padx=18,pady=15)
        years=sorted(set(get_available_years_for('sales')+get_available_years_for('purchases')), reverse=True); year_cb['values']=years
        if years: year_cb.set(years[0]); month_cb['values']=sorted(set(get_available_months_for(years[0],'sales')+get_available_months_for(years[0],'purchases')), key=lambda m: month_number_from_name(m) or 0)
        if month_cb['values']: month_cb.set(month_cb['values'][0])
        def update_summary(event=None):
            year=year_cb.get() or None; month=month_cb.get() or None
            sales=invoice_rows_filtered(year, month); purchases=purchase_rows_filtered(year, month)
            sales_taxable=sum(money(r[4]) for r in sales); sales_total=sum(money(r[10]) for r in sales)
            out_gst=sum(money(r[7])+money(r[8])+money(r[9]) for r in sales)
            purchase_taxable=sum(money(r[4]) for r in purchases); purchase_total=sum(money(r[8]) for r in purchases)
            in_gst=sum(money(r[5])+money(r[6])+money(r[7]) for r in purchases)
            profit=sales_taxable-purchase_taxable
            status='PROFIT' if profit>=0 else 'LOSS'
            summary.config(text=(
                f'MONTHLY BUSINESS SUMMARY\n'
                f'Month: {month or "ALL"} {year or "ALL"}\n\n'
                f'--------------- SALES ---------------\n'
                f'Total Bills                 : {len(sales)}\n'
                f'Total Sales Amount          : ₹{sales_total:.2f}\n'
                f'GST Collected from Customers: ₹{out_gst:.2f}\n\n'
                f'------------- PURCHASES -------------\n'
                f'Total Purchase Records      : {len(purchases)}\n'
                f'Total Purchase Amount       : ₹{purchase_total:.2f}\n'
                f'GST Paid on Purchases       : ₹{in_gst:.2f}\n\n'
                f'-------------- FINAL GST ------------\n'
                f'GST To Pay = GST Collected - GST Paid\n'
                f'Final GST To Pay            : ₹{(out_gst-in_gst):.2f}\n\n'
                f'---------- ESTIMATED EARNINGS -------\n'
                f'{status} / Estimated Earnings: ₹{abs(profit):.2f}'
            ))
        def year_changed(event=None):
            months=sorted(set(get_available_months_for(year_cb.get(),'sales')+get_available_months_for(year_cb.get(),'purchases')), key=lambda m: month_number_from_name(m) or 0)
            month_cb['values']=months; month_cb.set(months[0] if months else ''); update_summary()
        def do_report():
            try:
                # Saved report now focuses summary. Detail files remain in GST Export.
                path=export_report_filtered(year_cb.get() or None, month_cb.get() or None)
                messagebox.showinfo('Report Saved',f'Report saved:\n{path}')
                webbrowser.open(str(path))
            except Exception as e: messagebox.showerror('Report Error',str(e))
        year_cb.bind('<<ComboboxSelected>>',year_changed); month_cb.bind('<<ComboboxSelected>>',update_summary); update_summary()
        ttk.Button(box,text='Save Monthly Summary Report',command=do_report).grid(row=3,column=0,columnspan=2,sticky='w',padx=18,pady=15)

    def settings(self,main):
        self.wipe(main)
        tk.Label(main,text='Settings',font=('Segoe UI',22,'bold'),bg='#f8fafc',fg='#0f172a').pack(anchor='w',padx=25,pady=(12,8))

        # Scrollable settings area: fixes small laptop screens and .exe window sizes.
        outer=tk.Frame(main,bg='#f8fafc')
        outer.pack(fill='both',expand=True,padx=25,pady=(0,12))
        canvas=tk.Canvas(outer,bg='#f8fafc',highlightthickness=0)
        vbar=ttk.Scrollbar(outer,orient='vertical',command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side='right',fill='y')
        canvas.pack(side='left',fill='both',expand=True)
        content=tk.Frame(canvas,bg='#f8fafc')
        content_id=canvas.create_window((0,0),window=content,anchor='nw')

        def _on_frame_configure(event=None):
            canvas.configure(scrollregion=canvas.bbox('all'))
        def _on_canvas_configure(event):
            canvas.itemconfigure(content_id,width=event.width)
        content.bind('<Configure>',_on_frame_configure)
        canvas.bind('<Configure>',_on_canvas_configure)
        def _mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), 'units')
        canvas.bind_all('<MouseWheel>', _mousewheel)

        s=get_settings()
        box=tk.Frame(content,bg='white')
        box.pack(fill='x',pady=(0,12))
        keys=['name','gstin','address','email','phone','bank','account','ifsc','upi']
        ents={}
        for i,k in enumerate(keys):
            label = 'UPI ID (FOR QR CODE)' if k == 'upi' else k.upper()
            tk.Label(box,text=label,bg='white').grid(row=i,column=0,sticky='w',padx=12,pady=5)
            e=ttk.Entry(box,width=80)
            e.grid(row=i,column=1,padx=12,pady=5,sticky='ew')
            e.insert(0,s.get(k,''))
            ents[k]=e
        box.grid_columnconfigure(1,weight=1)
        def save():
            for k,e in ents.items():
                set_setting(k,e.get())
            messagebox.showinfo('Saved','Settings saved')
        ttk.Button(box,text='Save Settings',command=save).grid(row=len(keys),column=1,sticky='e',padx=12,pady=15)

        pathbox=tk.LabelFrame(content,text='File Saving Location',bg='white',font=('Segoe UI',11,'bold'))
        pathbox.pack(fill='x',pady=(0,12))
        tk.Label(pathbox,text='Master Records Folder',bg='white').grid(row=0,column=0,sticky='w',padx=12,pady=8)
        folder_var=tk.StringVar(value=s.get('records_folder', str(DATA)))
        folder_entry=ttk.Entry(pathbox,textvariable=folder_var,width=85)
        folder_entry.grid(row=0,column=1,padx=12,pady=8,sticky='ew')
        pathbox.grid_columnconfigure(1,weight=1)
        def browse_records_folder():
            folder=filedialog.askdirectory(title='Select Master Records Folder')
            if folder:
                folder_var.set(folder)
        def save_records_folder():
            selected=folder_var.get().strip() or str(DATA)
            set_setting('records_folder', selected)
            base=records_base()
            messagebox.showinfo('Saved', 'Records folder saved.\n\nFiles will save inside:\n' + str(base) + '\n\nSubfolders created: Invoices, GST Exports, Monthly Reports, Purchases, Backups')
        ttk.Button(pathbox,text='Browse',command=browse_records_folder).grid(row=0,column=2,padx=8,pady=8)
        ttk.Button(pathbox,text='Save Records Folder',command=save_records_folder).grid(row=1,column=1,sticky='e',padx=12,pady=(0,10))
        def manual_backup():
            path=create_backup('manual')
            if path: messagebox.showinfo('Backup Created', f'Backup saved:\n{path}')
            else: messagebox.showerror('Backup Error','Could not create backup.')
        ttk.Button(pathbox,text='Create Backup Now',command=manual_backup).grid(row=1,column=2,sticky='e',padx=8,pady=(0,10))
        def restore_backup_button():
            file=filedialog.askopenfilename(title='Select DEMO Backup ZIP', filetypes=[('ZIP backup','*.zip'),('All files','*.*')])
            if not file:
                return
            msg = ('Restore this backup?\n\n'
                   'This will restore database, customer/rate files, invoices, GST exports, purchases and monthly reports.\n'
                   'A safety backup of current data will be created first.\n\n'
                   'Close any open Word/Excel files before continuing.')
            if not messagebox.askyesno('Restore Backup', msg):
                return
            try:
                safety, restored = restore_backup_from_zip(file)
                messagebox.showinfo('Restore Complete', 'Backup restored successfully.\n\nRestored: ' + ', '.join(restored) + '\n\nSafety backup created at:\n' + str(safety) + '\n\nPlease close and reopen the app now.')
            except Exception as e:
                messagebox.showerror('Restore Failed', str(e))
        ttk.Button(pathbox,text='Restore Backup',command=restore_backup_button).grid(row=2,column=2,sticky='e',padx=8,pady=(0,10))
        tk.Label(pathbox,text='Tip: Backup now includes database, customer/rate files, invoices, GST exports, purchases and monthly reports.',bg='white',fg='#475569').grid(row=3,column=0,columnspan=3,sticky='w',padx=12,pady=(0,8))
        tk.Label(pathbox,text='Choose a safe folder like D:/DEMO Records. This also works in the .exe version without Python.',bg='white',fg='#475569').grid(row=4,column=0,columnspan=3,sticky='w',padx=12,pady=(0,8))

        passbox=tk.LabelFrame(content,text='Change Password',bg='white',font=('Segoe UI',11,'bold'))
        passbox.pack(fill='x',pady=(0,20))
        tk.Label(passbox,text='Old Password',bg='white').grid(row=0,column=0,sticky='w',padx=12,pady=6)
        old=ttk.Entry(passbox,show='*',width=35)
        old.grid(row=0,column=1,padx=12,pady=6,sticky='w')
        tk.Label(passbox,text='New Password',bg='white').grid(row=1,column=0,sticky='w',padx=12,pady=6)
        new=ttk.Entry(passbox,show='*',width=35)
        new.grid(row=1,column=1,padx=12,pady=6,sticky='w')
        tk.Label(passbox,text='Confirm New Password',bg='white').grid(row=2,column=0,sticky='w',padx=12,pady=6)
        confirm=ttk.Entry(passbox,show='*',width=35)
        confirm.grid(row=2,column=1,padx=12,pady=6,sticky='w')
        def change_pw():
            s=get_settings(); current=s.get('password_hash') or s.get('password', DEFAULT_PASSWORD)
            if not check_password(old.get(), current):
                messagebox.showerror('Wrong Password','Old password is incorrect')
                return
            if len(new.get().strip()) < 6 or new.get()!=confirm.get():
                messagebox.showerror('Invalid','New password must be at least 6 characters and confirmation must match')
                return
            set_setting('password_hash', hash_password(new.get()))
            create_backup('after_password_change')
            messagebox.showinfo('Changed','Password changed successfully')
            old.delete(0,'end'); new.delete(0,'end'); confirm.delete(0,'end')
        ttk.Button(passbox,text='Change Password',command=change_pw).grid(row=3,column=1,sticky='e',padx=12,pady=12)
        tk.Label(passbox,text='You can scroll this Settings page if your screen is small.',bg='white',fg='#64748b').grid(row=4,column=0,columnspan=2,sticky='w',padx=12,pady=(0,10))
        _on_frame_configure()

    def open_customers(self):
        ensure_master_files(); webbrowser.open(str(MASTER/'customers.xlsx'))
    def open_rates(self):
        ensure_master_files(); webbrowser.open(str(MASTER/'rate_list.xlsx'))

if __name__=='__main__': App().mainloop()
